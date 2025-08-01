# /scripts/merge_files.py (Corrected)

# =============================================================================
# SECTION 1: DEPENDENCIES
# =============================================================================
import os
import sys
import json
from docx import Document
from pypdf import PdfMerger
import pandas as pd
from logger import setup_logger
from ai_utils import get_ai_prettyname, get_ai_summary

logger = setup_logger('file_merger')

# =============================================================================
# SECTION 2: MERGE LOGIC FOR SPECIFIC FILE TYPES
# =============================================================================

def merge_pdfs(input_paths, output_path):
    """Merges multiple PDF documents into a single PDF."""
    merger = PdfMerger()
    try:
        for pdf_path in input_paths:
            logger.info(f"Appending PDF: {os.path.basename(pdf_path)}")
            merger.append(pdf_path)
        logger.info(f"Writing merged PDF to {output_path}")
        merger.write(output_path)
    finally:
        merger.close()

def merge_docx(input_paths, output_path):
    """Merges multiple DOCX documents into a single DOCX."""
    if not input_paths:
        return
    merged_document = Document(input_paths[0])
    for docx_path in input_paths[1:]:
        merged_document.add_page_break()
        logger.info(f"Appending DOCX: {os.path.basename(docx_path)}")
        sub_doc = Document(docx_path)
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    logger.info(f"Writing merged DOCX to {output_path}")
    merged_document.save(output_path)

def merge_text_based(input_paths, output_path, file_type):
    """Merges plain text files (TXT, MD, CSV)."""
    logger.info(f"Starting merge for {file_type} files.")
    with open(output_path, 'w', encoding='utf-8') as outfile:
        for file_path in input_paths:
            logger.info(f"Appending file: {os.path.basename(file_path)}")
            with open(file_path, 'r', encoding='utf-8') as infile:
                outfile.write(infile.read())
                if file_type in ['.txt', '.md']:
                    outfile.write(f'\n\n--- Merged from {os.path.basename(file_path)} ---\n\n')
                else:
                    outfile.write('\n')
    logger.info(f"Finished merging {file_type} files to {output_path}")

def merge_xlsx(input_paths, output_path):
    """Merges multiple XLSX files into a single XLSX with multiple sheets."""
    logger.info("Starting merge for XLSX files.")
    with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
        for xlsx_path in input_paths:
            try:
                logger.info(f"Reading sheet from: {os.path.basename(xlsx_path)}")
                df = pd.read_excel(xlsx_path)
                sheet_name = os.path.splitext(os.path.basename(xlsx_path))[0]
                safe_sheet_name = ''.join(e for e in sheet_name if e.isalnum())[:31]
                df.to_excel(writer, sheet_name=safe_sheet_name, index=False)
            except Exception as e:
                logger.error(f"Could not process Excel file {xlsx_path}: {e}")
    logger.info(f"Finished merging XLSX files to {output_path}")

# =============================================================================
# SECTION 3: AI-POWERED "MERGE WITH CONTENTS" LOGIC
# =============================================================================

def add_contents_and_covers(input_paths, output_path, first_ext):
    """Creates a merged document with a table of contents and AI-generated summaries."""
    if first_ext not in ['.docx', '.txt', '.md']:
        logger.warning(f"'With Contents' mode not supported for {first_ext}, performing straight merge instead.")
        return False, 0 # Return False for success and 0 for tokens

    contents = "Table of Contents\n\n"
    covers = []
    total_tokens_used = 0 # Initialize token counter for this operation

    for file_path in input_paths:
        # Capture both the name and the tokens used
        pretty_title, name_tokens = get_ai_prettyname(file_path)
        total_tokens_used += name_tokens
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content_for_summary = f.read()
        except Exception:
            content_for_summary = "Cannot read file content."

        # Capture both the summary and the tokens used
        summary, summary_tokens = get_ai_summary(content_for_summary)
        total_tokens_used += summary_tokens
        
        covers.append({'title': pretty_title, 'summary': summary})
        contents += f"- {pretty_title} (Original: {os.path.basename(file_path)})\n"


    if first_ext == '.docx':
        merged_doc = Document()
        merged_doc.add_heading('Merged Document', level=0)
        merged_doc.add_paragraph(contents)
        for i, file_path in enumerate(input_paths):
            merged_doc.add_page_break()
            merged_doc.add_heading(covers[i]['title'], level=1)
            merged_doc.add_paragraph(covers[i]['summary'])
            merged_doc.add_page_break()
            sub_doc = Document(file_path)
            for element in sub_doc.element.body:
                merged_doc.element.body.append(element)
        merged_doc.save(output_path)
    else: # For .txt and .md
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# Merged Document\n\n{contents}\n\n")
            for i, file_path in enumerate(input_paths):
                f.write('---\n\n')
                f.write(f"## {covers[i]['title']}\n\n")
                f.write(f"**Summary:** {covers[i]['summary']}\n\n")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as infile:
                    f.write(infile.read())
                f.write('\n\n')
    return True, total_tokens_used

# =============================================================================
# SECTION 4: MAIN FUNCTION
# =============================================================================

def merge_files(api_key, payload, output_dir):
    """Main function to handle file merging based on payload."""
    input_paths = payload.get('files', [])
    mode = payload.get('mode', 'straight')
    tokens_used = 0 # Initialize tokens for the entire operation
    
    if not input_paths:
        raise ValueError("No input files were provided for merging.")

    first_ext = os.path.splitext(input_paths[0])[1].lower()
    if not all(os.path.splitext(p)[1].lower() == first_ext for p in input_paths):
        raise ValueError("All files selected for merging must be of the same type.")

    output_filename = f"merged_document{first_ext}"
    output_path = os.path.join(output_dir, output_filename)

    message = ""
    if mode == 'with_contents':
        success, tokens_from_ai = add_contents_and_covers(input_paths, output_path, first_ext)
        tokens_used += tokens_from_ai
        if success:
            message = f"Successfully merged {len(input_paths)} files with AI-generated contents."
            # Return a JSON string that includes the token count
            return json.dumps({"message": message, "tokens_used": tokens_used})

    # Fallback to straight merge if 'with_contents' is not supported or fails
    if first_ext == '.pdf':
        merge_pdfs(input_paths, output_path)
    elif first_ext == '.docx':
        merge_docx(input_paths, output_path)
    elif first_ext in ['.md', '.txt', '.csv']:
        merge_text_based(input_paths, output_path, first_ext)
    elif first_ext == '.xlsx':
        merge_xlsx(input_paths, output_path)
    else:
        raise ValueError(f"Unsupported file type for merging: {first_ext}")
        
    message = f"Successfully merged {len(input_paths)} files."
    # Return a JSON string for standard merges as well, with 0 tokens
    return json.dumps({"message": message, "tokens_used": 0})